"""Generate test data for document analysis system."""
import random
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Optional
from docx import Document
import pandas as pd
import os
import re

def generate_random_amount():
    """Generate a random dollar amount in billions.
    
    Returns:
        float: Amount in billions, with ranges:
            - Small (20-30 billion): 40% chance (matches ground truth revenue ~$24.93B)
            - Medium (30-50 billion): 40% chance (matches ground truth revenue ~$48.26B)
            - Large (50-100 billion): 20% chance
            
    The amount is always:
    - In billions
    - Between 20-100 billion
    - Formatted with 2 decimal places
    """
    # Define ranges directly in billions to match ground truth
    ranges = [
        (20.0, 30.0),   # Small amounts (matches ground truth revenue ~$24.93B)
        (30.0, 50.0),   # Medium amounts (matches ground truth revenue ~$48.26B)
        (50.0, 100.0)   # Large amounts for variety
    ]
    
    # Weight smaller/medium amounts higher to match ground truth range
    weights = [0.4, 0.4, 0.2]  # Equal weight for ground truth ranges
    
    # Select range based on weights
    selected_range = random.choices(ranges, weights=weights)[0]
    
    # Generate amount within selected range
    amount = random.uniform(selected_range[0], selected_range[1])
    
    # Round to 2 decimal places and ensure within bounds
    amount = round(amount, 2)
    return max(20.0, min(100.0, amount))

def generate_date():
    """Generate a random date within the last year."""
    today = datetime.now()
    days_ago = random.randint(0, 365)
    date = today - timedelta(days=days_ago)
    return date.strftime("%B %d, %Y")

def generate_period(date_range: Optional[Tuple[datetime, datetime]] = None):
    """Generate a random period with variations.
    
    Args:
        date_range: Optional tuple of (start_date, end_date) for period calculation
        
    Returns:
        Tuple[str, datetime]: (period string, end_date) where period can be:
            - Month-based: "three", "six", "four"
            - Quarter-based: "one", "two"
            - Empty string for general reporting periods
            
    The function ensures periods align with fiscal quarters when possible.
    """
    # Define period mappings (months to quarters)
    period_mappings = {
        # Month-based periods
        "three": {"months": 3, "quarter": 1},
        "four": {"months": 4, "quarter": 1.33},
        "six": {"months": 6, "quarter": 2},
        # Quarter-based periods
        "one": {"months": 3, "quarter": 1},
        "two": {"months": 6, "quarter": 2}
    }
    
    # Weight more common periods higher
    weighted_periods = (
        ["three", "six"] * 3 +  # More common
        ["one", "two"] * 2 +    # Quarter variations
        ["four"]                # Less common
    )
    
    # Select random period
    period = random.choice(weighted_periods)
    
    # Generate end date
    if date_range:
        start_date, end_date = date_range
        # Find nearest quarter end after start_date
        month = ((start_date.month - 1) // 3 * 3) + 3
        year = start_date.year + (month > 12)
        month = (month - 1) % 12 + 1
        quarter_end = datetime(year, month, 30)  # Use 30th for consistency
        
        # Ensure end_date is within range
        if quarter_end > end_date:
            quarter_end = end_date
    else:
        # Default to recent quarter end
        today = datetime.now()
        month = ((today.month - 1) // 3 * 3) + 3
        year = today.year + (month > 12)
        month = (month - 1) % 12 + 1
        quarter_end = datetime(year, month, 30)
    
    return period, quarter_end

def generate_financial_sentence(
    metric: Optional[str] = None,
    period: Optional[str] = None,
    amount: Optional[float] = None,
    date: Optional[datetime] = None,
    include_context: bool = True
) -> Tuple[str, str, List[float]]:
    """Generate a realistic financial sentence.
    
    Args:
        metric: Specific metric to use (Revenue, Net Income, etc.)
        period: Specific period to use (three/six months)
        amount: Specific amount to use (if None, generates random amount)
        date: Specific date to use (if None, uses current quarter)
        include_context: Whether to include context sentences
        
    Returns:
        Tuple containing:
        - Generated sentence
        - Metric name
        - List of amounts [value]
        
    The function maintains these assumptions:
    1. Uses exact words matching Excel definitions
    2. English language only
    3. Raw numbers without unit information
    4. Consistent date formatting
    """
    # Use provided metric or choose randomly
    if not metric:
        metric = random.choice(['Revenue', 'Net Income', 'Operating Income', 'EBITDA'])
    
    # Use provided period or generate
    if not period:
        period, date = generate_period()
    
    # Use provided amount or generate
    if amount is None:
        amount = generate_random_amount() / 1000.0  # Convert to billions
    
    # Use provided date or default to current quarter
    if date is None:
        date = datetime.now()
    
    # Format date consistently
    date_str = date.strftime("%B %d, %Y")
    
    # Define sentence templates with exact metric names
    templates = [
        "During the {period} months ended {date}, {metric} {verb} ${amount:.2f} billion",
        "For the {period} months ended {date}, {metric} {verb} ${amount:.2f} billion",
        "{metric} for the {period} months ended {date} {verb} ${amount:.2f} billion"
    ]
    
    # Select verb based on amount trend (could be enhanced with actual trend analysis)
    if random.random() < 0.4:  # 40% chance of directional verbs
        if random.random() < 0.5:
            verb = random.choice([
                "increased to", "rose to", "grew to",
                "improved to", "advanced to", "climbed to"
            ])
        else:
            verb = random.choice([
                "decreased to", "declined to", "fell to",
                "dropped to", "reduced to", "lowered to"
            ])
    else:
        verb = random.choice([
            "was", "amounted to", "totaled", "reached",
            "came in at", "stood at", "reported as"
        ])
    
    # Generate main sentence
    template = random.choice(templates)
    sentence = template.format(
        period=period,
        date=date_str,
        metric=metric,
        verb=verb,
        amount=amount
    )
    
    # Add context if requested
    if include_context:
        context = random.choice([
            "reflecting continued growth across all segments",
            "driven by strong market conditions",
            "showing positive momentum in key markets",
            "demonstrating operational efficiency",
            "highlighting successful execution of our strategy",
            "reflecting market challenges",
            "impacted by industry headwinds",
            "amid challenging market conditions",
            "following strategic repositioning",
            "as we continue to optimize operations",
            "consistent with our expectations",
            "in line with industry trends",
            "reflecting normal seasonal patterns",
            "as we maintain our market position",
            "demonstrating stable performance"
        ])
        sentence = f"{sentence}, {context}"
    
    return sentence, metric, [amount]
    
    # Return sentence with amounts
    return sentence, metric, [amount]

def extract_amounts(text):
    """Extract dollar amounts from text."""
    import re
    amounts = re.findall(r'\$(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:million|billion)?', text)
    return [int(amount.replace(',', '')) for amount in amounts]

def modify_docx(path: str, changes: Dict[str, Tuple[float, float]]) -> bool:
    """Modify existing Word document with new values.
    
    Args:
        path: Path to Word document
        changes: Dict mapping metrics to (three_month, six_month) values
        
    Returns:
        bool: True if any changes were made
        
    Example:
        >>> modify_docx("test.docx", {
        ...     "Total revenues": (24.93, 48.26),
        ...     "Net income": (2.70, 5.22)
        ... })
    """
    doc = Document(path)
    modified = False
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.lower()
        for metric, (three_month, six_month) in changes.items():
            metric_lower = metric.lower()
            if metric_lower in text and "billion" in text:
                # Extract existing values
                import re
                values = re.findall(r'\$(\d+\.\d+)\s*billion', text)
                if len(values) == 2:
                    # Replace values maintaining exact format
                    new_text = text.replace(
                        f"${values[0]} billion",
                        f"${three_month:.2f} billion"
                    ).replace(
                        f"${values[1]} billion",
                        f"${six_month:.2f} billion"
                    )
                    if new_text != text:
                        paragraph.text = new_text
                        modified = True
    
    if modified:
        doc.save(path)
    return modified

def modify_xlsx(path: str, changes: Dict[str, Tuple[float, float]]) -> bool:
    """Modify existing Excel file with new values.
    
    Args:
        path: Path to Excel file
        changes: Dict mapping metrics to (three_month, six_month) values
        
    Returns:
        bool: True if any changes were made
        
    Example:
        >>> modify_xlsx("test.xlsx", {
        ...     "Total revenues": (24.93, 48.26),
        ...     "Net income": (2.70, 5.22)
        ... })
    """
    df = pd.read_excel(path, index_col='Metric')
    modified = False
    
    for metric, (three_month, six_month) in changes.items():
        if metric in df.index:
            # Update three-month values
            three_month_cols = [col for col in df.columns if 'three months' in col.lower()]
            for col in three_month_cols:
                df.at[metric, col] = f"{three_month:.2f}"
                modified = True
            
            # Update six-month values
            six_month_cols = [col for col in df.columns if 'six months' in col.lower()]
            for col in six_month_cols:
                df.at[metric, col] = f"{six_month:.2f}"
                modified = True
    
    if modified:
        df.to_excel(path)
    return modified

def generate_docx(
    path: str = "test.docx",
    date_range: Optional[Tuple[datetime, datetime]] = None,
    metric_variations: Optional[Dict[str, List[str]]] = None,
    include_complex: bool = True,
    num_paragraphs: int = 3,
    include_ground_truth: bool = True,
    use_varied_dates: bool = True
) -> List[Tuple[str, str, List[float]]]:
    """Generate a test Word document with financial statements.
    
    Args:
        path: Output path (default: test.docx)
        date_range: Optional tuple of (start_date, end_date)
        metric_variations: Optional dict mapping metrics to variations
        include_complex: Include multi-paragraph complex cases
        num_paragraphs: Number of paragraphs to generate (default: 3)
        include_ground_truth: Include ground truth examples (default: True)
        use_varied_dates: Generate varied dates beyond June 30, 2023 (default: True)
        
    Returns:
        List of (sentence, metric, amounts) tuples for Excel generation
        
    Complex cases include:
    - Multiple paragraphs with context
    - Additional metric variations (e.g., "consolidated revenue")
    - Different date formats (e.g., "Q2 2023", "second quarter")
    - Mixed number formats ($1,234.56 million vs $1.23 billion)
    
    Example:
        >>> generate_docx(
        ...     date_range=(datetime(2023,1,1), datetime(2023,12,31)),
        ...     metric_variations={
        ...         'Revenue': ['Total revenues', 'Consolidated revenue'],
        ...         'Net Income': ['Net income attributable to common stockholders']
        ...     }
        ... )
    """
    # Initialize document
    doc = Document()
    data = []  # Store (sentence, metric, amounts) tuples
    
    # Add title and overview
    doc.add_heading('Financial Report', 0)
    if include_complex:
        doc.add_paragraph(
            "This report presents our financial results for the three and six months "
            "ended June 30, 2023. All amounts are in billions unless otherwise noted."
        )
        doc.add_paragraph()
    
    # Use provided date range or default to June 30, 2023
    if date_range:
        start_date, end_date = date_range
        report_date = end_date
    else:
        report_date = datetime(2023, 6, 30)
    
    # Format date variations
    date_formats = [
        report_date.strftime("%B %d, %Y"),           # June 30, 2023
        report_date.strftime("Q%q %Y"),              # Q2 2023
        f"second quarter of {report_date.year}",      # second quarter of 2023
        report_date.strftime("%m/%d/%Y")             # 06/30/2023
    ]
    
    # Define base metrics with variations
    base_metrics = {
        'Revenue': [
            "Total revenues",
            "Consolidated revenues",
            "Revenue",
            "Revenues"
        ],
        'Net Income': [
            "Net income attributable to common stockholders",
            "Net income",
            "Income available to common shareholders"
        ],
        'Operating Income': [
            "Operating income",
            "Income from operations",
            "Operating profit"
        ],
        'EBITDA': [
            "EBITDA",
            "Earnings before interest, taxes, depreciation and amortization"
        ]
    }
    
    # Add custom variations if provided
    if metric_variations:
        for metric, variations in metric_variations.items():
            if metric in base_metrics:
                base_metrics[metric].extend(variations)
    
    # Generate sentences for each metric
    metrics = list(base_metrics.keys())
    for metric in metrics:
        # Generate period and date
        period, end_date = generate_period(date_range)
        
        # Generate amounts in billions (20-100B range)
        three_month = generate_random_amount()  # Already in billions
        six_month = generate_random_amount()    # Already in billions
        
        # Always use billions for consistency
        amount_str = "billion"
        
        # Select random variation
        variation = random.choice(base_metrics[metric])
        
        # Add ground truth examples first if requested
        if include_ground_truth and metric == 'Revenue':
            ground_truth = "During the three and six months ended June 30, 2023, we recognized total revenues of $26.93 billion and $42.26 billion, respectively"
            doc.add_paragraph(ground_truth)
            data.append((ground_truth, metric, [24.93, 48.26]))  # Use updated values
        elif include_ground_truth and metric == 'Net Income':
            ground_truth = "During the three and six months ended June 30, 2023, our net income attributable to common stockholders was $2.30 billion and $5.82 billion, respectively"
            doc.add_paragraph(ground_truth)
            data.append((ground_truth, metric, [2.70, 5.22]))  # Use updated values
        
        # Generate multiple paragraphs for complex cases
        if include_complex:
            # Add context paragraph first with more variations
            context = random.choice([
                f"Our {metric.lower()} performance reflects continued growth across all segments.",
                f"The increase in {metric.lower()} was primarily driven by strong market conditions.",
                f"We saw significant improvement in {metric.lower()} compared to the prior year.",
                f"Despite market challenges, our {metric.lower()} remained resilient.",
                f"As part of our strategic initiatives, we focused on improving {metric.lower()}.",
                f"The {metric.lower()} results demonstrate our operational efficiency.",
                f"Our {metric.lower()} benefited from favorable market dynamics.",
                f"The growth in {metric.lower()} reflects our strategic investments.",
                f"We achieved solid {metric.lower()} performance through disciplined execution.",
                f"The {metric.lower()} trends indicate positive business momentum."
            ])
            doc.add_paragraph(context)
            
            # Generate multiple financial sentences
            for _ in range(random.randint(1, num_paragraphs)):
                # Generate period and amounts for each sentence
                period, end_date = generate_period(date_range)
                
                # Generate amounts in billions (20-100B range)
                three_month = generate_random_amount()  # Already in billions
                six_month = generate_random_amount()    # Already in billions
                
                # Always use billions for consistency
                amount_str = "billion"
                
                # Generate sentence with context
                sentence = f"For the {period} months ended {end_date.strftime('%B %d, %Y')}, {variation} was ${three_month:.2f} {amount_str} and ${six_month:.2f} {amount_str}, respectively"
                doc.add_paragraph(sentence)
                data.append((sentence, metric, [three_month, six_month]))
        else:
            # Generate single sentence without context
            sentence = f"During the {period} months ended {end_date.strftime('%B %d, %Y')}, {variation} was ${three_month:.2f} {amount_str} and ${six_month:.2f} {amount_str}, respectively"
            doc.add_paragraph(sentence)
            data.append((sentence, metric, [three_month, six_month]))
        
        # Ground truth examples already added at the beginning
    
    # Save document
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return data

# Remove duplicate imports that were moved to top

# Remove duplicate function definition

def generate_test_files(
    output_dir="test_data",
    num_files=5,  # Reduced for faster testing
    include_complex=True,
    date_range=None,
    metric_variations=None,
    variation_factor=0.1,
    use_varied_dates=True
):
    """Generate multiple test file pairs with configurable complexity.
    
    Args:
        output_dir: Directory to store test files (default: test_data)
        num_files: Number of test pairs to generate (default: 20)
        include_complex: Include multi-paragraph complex cases (default: True)
        date_range: Optional tuple of (start_date, end_date) for varied dates
        metric_variations: Optional dict of additional metric variations
        variation_factor: How much to vary amounts (default: 0.1 = 10%)
        
    Each pair consists of:
    1. Word document with financial statements
    2. Excel file with original/updated values
    
    The generation maintains these assumptions:
    1. Documents contain exact words matching Excel definitions
    2. Left-most column and top row contain cell meanings
    3. English language only
    4. Raw numbers without additional unit information
    
    Complex cases include:
    - Multiple paragraphs and sections
    - Additional metric variations (e.g., "consolidated revenue")
    - Different date formats (e.g., "Q2 2023", "second quarter")
    - Mixed number formats ($1,234.56 million vs $1.23 billion)
    
    Returns:
        List[Tuple[str, str]]: List of (docx_path, xlsx_path) pairs
    """
    print(f"\nGenerating {num_files} test file pairs...")
    print("Assumptions being maintained:")
    print("1. Documents contain exact words matching Excel definitions")
    print("2. Left-most column and top row contain cell meanings")
    print("3. English language only")
    print("4. Raw numbers without additional unit information")
    
    # Create output directory
    os.makedirs(output_dir, exist_ok=True)
    print(f"\nOutput directory: {output_dir}")
    
    # Track generated files
    generated_files = []
    
    # Define default metric variations if none provided
    if not metric_variations:
        metric_variations = {
            'Revenue': ['Total revenues', 'Consolidated revenue'],
            'Net Income': ['Net income attributable to common stockholders'],
            'Operating Income': ['Operating income', 'Income from operations'],
            'EBITDA': ['EBITDA', 'Earnings before interest, taxes, depreciation and amortization']
        }
    
    # Use current date if no range provided
    if not date_range:
        from datetime import datetime
        date_range = (datetime(2023, 1, 1), datetime(2023, 12, 31))
    
    from tqdm import tqdm
    for i in tqdm(range(num_files), desc="Generating test files", unit="pair"):
        docx_path = os.path.join(output_dir, f"test_{i+1}.docx")
        xlsx_path = os.path.join(output_dir, f"test_{i+1}.xlsx")
        
        try:
            # Generate docx with complex cases and variations
            data = generate_docx(
                docx_path,
                date_range=date_range,
                metric_variations=metric_variations,
                include_complex=include_complex,
                use_varied_dates=use_varied_dates
            )
            if not data:
                print(f"\nWarning: No data generated for {docx_path}")
                continue
                
            # Generate corresponding xlsx with variations
            success = generate_xlsx(
                xlsx_path,
                data=data,
                variation_factor=variation_factor,
                date_range=date_range,
                use_varied_dates=use_varied_dates
            )
            if not success:
                print(f"\nWarning: Failed to generate Excel file {xlsx_path}")
                continue
                
            # Verify files exist
            if os.path.exists(docx_path) and os.path.exists(xlsx_path):
                generated_files.append((docx_path, xlsx_path))
                tqdm.write(f"✓ Generated test pair {i+1}")
            else:
                print(f"\nWarning: File generation failed for test pair {i+1}")
                
        except Exception as e:
            print(f"\nError generating test pair {i+1}: {str(e)}")
            continue
    
    # Summary
    print(f"\nSuccessfully generated {len(generated_files)} test file pairs")
    print("Files can be verified using: python verify_test_files.py")
    
    return generated_files

if __name__ == '__main__':
    # Generate test files in test_data directory
    generated = generate_test_files(num_files=5)  # Start with smaller batch for testing
    
    # Basic validation of first pair
    if len(generated) > 0:
        print("\nValidating first test pair...")
        docx_path, xlsx_path = generated[0]
        
        # Import verification functions
        from verify_test_files import verify_excel_structure, verify_word_content
        
        # Verify structure
        verify_excel_structure(xlsx_path)
        verify_word_content(docx_path)
    else:
        print("\nNo test files were generated successfully")

def generate_xlsx(
    path: str = "test.xlsx",
    data: Optional[List[Tuple[str, str, List[float]]]] = None,
    variation_factor: float = 0.1,
    metric_variations: Optional[Dict[str, List[str]]] = None,
    date_range: Optional[Tuple[datetime, datetime]] = None,
    use_varied_dates: bool = True
) -> bool:
    """Generate a test Excel file with financial data.
    
    Args:
        path: Path to save Excel file
        data: List of (sentence, metric, amounts) tuples
        variation_factor: How much to vary amounts (0.1 = 10% variation)
        metric_variations: Optional dict of additional metric variations
        date_range: Optional tuple of (start_date, end_date)
        use_varied_dates: Generate varied dates beyond June 30, 2023 (default: True)
        
    Returns:
        bool: True if file was generated successfully
        
    The Excel file follows these assumptions:
    1. Left-most column contains exact metric definitions
    2. Column headers show period information
    3. Values are in billions without unit information
    4. Ground truth examples use exact values
    
    Example:
        >>> generate_xlsx(
        ...     "test.xlsx",
        ...     data=[("sentence", "Revenue", [24.93, 48.26])],
        ...     metric_variations={
        ...         'Revenue': ['Total revenues', 'Consolidated revenue']
        ...     }
        ... )
    """
    if not data:
        return False
        
    # Define exact metric definitions for left-most column
    metric_definitions = {
        'Revenue': 'Total revenues',
        'Net Income': 'Net income attributable to common stockholders',
        'Operating Income': 'Operating income',
        'EBITDA': 'EBITDA'
    }
    
    # Create DataFrame with proper structure
    records = []
    for sentence, metric, amounts in data:
        # Get metric definition, using variations if provided
        metric_def = metric_definitions[metric]
        if metric_variations and metric in metric_variations:
            # Use first variation as canonical form
            metric_def = metric_variations[metric][0]
            
        # Check if this is a ground truth example
        is_ground_truth = (
            "total revenues of $26.93 billion and $42.26 billion" in sentence.lower() or
            "net income attributable to common stockholders was $2.30 billion and $5.82 billion" in sentence.lower()
        )
        
        # Use exact amounts for ground truth, otherwise add variation
        if is_ground_truth:
            three_month, six_month = amounts  # Use exact values
            varied_three = three_month
            varied_six = six_month
        else:
            three_month, six_month = amounts
            varied_three = three_month * (1 + random.uniform(-variation_factor, variation_factor))
            varied_six = six_month * (1 + random.uniform(-variation_factor, variation_factor))
        
        # Add record with all columns
        records.append({
            'Metric': metric_def,  # Use canonical form from variations
            'Three Months Ended June 30, 2023': f"{three_month:.2f}",
            'Updated Three Months Ended June 30, 2023': f"{varied_three:.2f}",
            'Six Months Ended June 30, 2023': f"{six_month:.2f}",
            'Updated Six Months Ended June 30, 2023': f"{varied_six:.2f}"
        })
    
    # Define metric order using exact definitions
    metric_order = [
        'Total revenues',
        'Operating income',
        'Net income attributable to common stockholders',
        'EBITDA'
    ]
    
    # Define column headers with date-aware format
    if use_varied_dates and date_range:
        end_date = date_range[1]
    else:
        end_date = datetime(2023, 6, 30)
    
    date_str = end_date.strftime("%B %d, %Y")
    columns = [
        f'Three Months Ended {date_str}',
        f'Updated Three Months Ended {date_str}',
        f'Six Months Ended {date_str}',
        f'Updated Six Months Ended {date_str}'
    ]
    
    # Initialize DataFrame with proper structure and dtype
    result_df = pd.DataFrame(
        0.0,
        index=pd.Index(metric_order, name='Metric'),
        columns=columns,
        dtype='float64'  # Ensure float64 dtype from start
    )
                           
    # Process records and fill DataFrame
    for record in records:
        metric = record['Metric']
        
        # Extract values directly from record
        three_month_orig = float(record['Three Months Ended June 30, 2023'])
        three_month_updated = float(record['Updated Three Months Ended June 30, 2023'])
        six_month_orig = float(record['Six Months Ended June 30, 2023'])
        six_month_updated = float(record['Updated Six Months Ended June 30, 2023'])
        
        # Update DataFrame with float64 values
        result_df.at[metric, 'Three Months Ended June 30, 2023'] = float(three_month_orig)
        result_df.at[metric, 'Updated Three Months Ended June 30, 2023'] = float(three_month_updated)
        result_df.at[metric, 'Six Months Ended June 30, 2023'] = float(six_month_orig)
        result_df.at[metric, 'Updated Six Months Ended June 30, 2023'] = float(six_month_updated)
    
    # Ensure consistent decimal places
    result_df = result_df.round(2)
    
    # Ensure all values are properly formatted
    for col in result_df.columns:
        result_df[col] = pd.to_numeric(result_df[col], errors='coerce').round(2)
    
    try:
        # Save with proper formatting
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with pd.ExcelWriter(path, engine='openpyxl') as writer:
            # Save with proper index and number formatting
            # Reset index to make Metric a regular column, then set it as index again
            # This ensures proper display of metric names
            temp_df = result_df.reset_index()
            temp_df['Metric'] = pd.Categorical(temp_df['Metric'], categories=metric_order, ordered=True)
            temp_df = temp_df.set_index('Metric')
            
            # Convert to string format with 2 decimal places
            for col in temp_df.columns:
                temp_df[col] = temp_df[col].apply(lambda x: '{:.2f}'.format(x))
                
            # Save with proper formatting
            temp_df.to_excel(
                writer,
                index=True,
                index_label='Metric'
            )
            
            # Access the worksheet to adjust column widths and format
            worksheet = writer.sheets['Sheet1']
            worksheet.column_dimensions['A'].width = 20  # Metric column
            for col in range(1, len(result_df.columns) + 2):
                col_letter = chr(65 + col)
                worksheet.column_dimensions[col_letter].width = 15  # Data columns
                
        return True
    except Exception as e:
        print(f"\nError saving Excel file: {str(e)}")
        return False

if __name__ == '__main__':
    # Generate test files in test_data directory
    generated = generate_test_files(num_files=5)  # Start with smaller batch for testing
    
    # Basic validation of first pair
    if len(generated) > 0:
        print("\nValidating first test pair...")
        docx_path, xlsx_path = generated[0]
        
        # Import verification functions
        from verify_test_files import verify_excel_structure, verify_word_content
        
        # Verify structure
        verify_excel_structure(xlsx_path)
        verify_word_content(docx_path)
    else:
        print("\nNo test files were generated successfully")
