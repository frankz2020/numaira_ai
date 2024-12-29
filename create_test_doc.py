from docx import Document

doc = Document()

# Add title
doc.add_paragraph('ITEM 2. MANAGEMENT\'S DISCUSSION AND ANALYSIS OF FINANCIAL CONDITION AND RESULTS OF OPERATIONS')

# Add overview
doc.add_paragraph('Overview\nOur mission is to provide innovative solutions in the technology sector. We continue to invest in research and development while expanding our market presence globally.')

# Add financial results
doc.add_paragraph('Financial Results\nDuring the three and six months ended September 30, 2023, our consolidated group total revenue and other income reached $18.45 billion and $35.67 billion, respectively. Our growth was driven by strong performance across all business segments.')

# Add net income
doc.add_paragraph('During the three and six months ended September 30, 2023, our net income available to common shareholders after non-controlling interests was $3.12 billion and $6.45 billion, respectively. This performance reflects our operational efficiency and market strategy.')

# Add segment performance
doc.add_paragraph('Segment Performance\nOur Enterprise Software Solutions and Services Revenue Stream generated $8.92 billion and $17.34 billion for the three and six months ended September 30, 2023, respectively. The Infrastructure and Cloud Computing Platform Services division contributed $5.67 billion and $10.89 billion for the same periods.')

# Add R&D expenses
doc.add_paragraph('Research Development and Technology Innovation expenses were $1.23 billion and $2.45 billion for the three and six months ended September 30, 2023, as we continued to invest in next-generation technologies.')

# Add operating expenses
doc.add_paragraph('Total Operating and Administrative Expenses Including Non-recurring Items were $2.78 billion and $5.43 billion for the three and six months ended September 30, 2023, reflecting our ongoing investments in market expansion and operational improvements.')

# Save the document
doc.save('test_data/test_financial_report.docx') 