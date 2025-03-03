import logging
from typing import Dict, List, Tuple, Any
from docx import Document
from docx.shared import Pt, RGBColor
from config import Config
from funnels.extract import read_docx  # Utility for extracting text from Word documents.
from funnels.excel_utils import excel_to_list  # Utility for converting Excel data to a list.
from funnels.llm_provider import get_llm_provider  # LLM integration for text updates.
from tqdm import tqdm  # Progress bar utility for loops.

logger = logging.getLogger(__name__)

def calculate_confidence(sentence: str, header: str, values: Dict[str, float], updated_text: str, num_matches: int) -> float:
    """Calculate confidence score for a match and update.
    
    Factors considered:
    1. Match quality (length ratio of match to sentence)
    2. Position of match in sentence (earlier is better)
    3. Successful value updates (ratio of values updated)
    4. Number of competing matches (penalizes multiple matches)
    
    Returns:
        Confidence score between 0.0 and 1.0
    """
    # Base confidence for exact match
    confidence = 0.9
    
    # Factor 1: Length ratio of match to sentence (0-0.15)
    # Longer matches relative to sentence length are more likely to be correct
    length_ratio = len(header) / len(sentence)
    # Scale ratio to max contribution of 0.15
    length_score = min(length_ratio * 0.5, 0.15)
    confidence += length_score
    
    # Factor 2: Position in sentence (0-0.15)
    # Earlier matches are more likely to be the main subject
    position = sentence.lower().find(header.lower()) / len(sentence)
    # Invert position (0 is best) and scale to max contribution of 0.15
    position_score = (1 - position) * 0.15
    confidence += position_score
    
    # Factor 3: Value update success (0-0.1)
    if updated_text != sentence:
        # Calculate ratio of values successfully updated
        updated_values = sum(1 for value in values.values() 
                           if str(value) in updated_text)
        update_ratio = updated_values / len(values)
        # Scale ratio to max contribution of 0.1
        update_score = update_ratio * 0.1
        confidence += update_score
    
    # Factor 4: Multiple match penalty
    # Each additional match reduces confidence by 0.1
    if num_matches > 1:
        penalty = 0.1 * (num_matches - 1)
        confidence -= penalty
    
    # Ensure confidence stays between 0 and 1
    return max(min(confidence, 1.0), 0.0)

def find_exact_matches(sentence: str, excel_data: List[Dict[str, Any]]) -> List[Tuple[str, Dict[str, float]]]:
    """Find exact matches between sentence and Excel data."""
    matches = []
    for row in excel_data:
        header = row["row_header"].lower()
        if header in sentence.lower():
            matches.append((header, row["values"]))
    return matches

def process_files_with_selected_data(docx_data: dict, excel_data: list, timeout: int = None) -> List[Tuple[str, str, float]]:
    """Process Word and Excel files to find and update matching content."""
    try:
        # Load input files
        print("Loading files...")
        sentences = docx_data.values()
        
        if not sentences:
            raise ValueError("No text found in the Word document")
            
        # Initialize LLM provider
        config = Config()
        llm_config = config.get_model_config()["llm"]
        llm = get_llm_provider(llm_config["provider"])
        
        # Process each sentence
        results = []
        print("\nProcessing sentences...")
        for sentence in tqdm(sentences, desc="Analyzing", unit="sentence"):
            # Find exact matches in sentence
            matches = find_exact_matches(sentence, excel_data)
            if not matches:
                continue
            
            # Use LLM to update sentence with new values
            prompt = f"""Update this sentence with the new values from Excel:

Sentence: {sentence}

Excel data:
{chr(10).join(f'- {header}: {values}' for header, values in matches)}

Rules:
1. Keep the same sentence structure
2. Only update the numerical values
3. Keep all formatting (e.g., "$", "billion")
4. Do not change any other text

Output only the updated sentence."""

            try:
                updated_text = llm.analyze_text(prompt, timeout)
                if isinstance(updated_text, dict):
                    updated_text = updated_text.get("analysis", sentence)
                
                if updated_text and updated_text != sentence:
                    # Calculate confidence for each match
                    match_confidences = [
                        calculate_confidence(sentence, header, values, updated_text, len(matches))
                        for header, values in matches
                    ]
                    # Use average confidence if multiple matches
                    confidence = sum(match_confidences) / len(match_confidences)
                    
                    results.append((
                        sentence,
                        updated_text,
                        confidence
                    ))
                    
            except Exception as e:
                logger.error(f"Error updating sentence: {str(e)}")
                continue
        
        return results
            
    except Exception as e:
        logger.error(f"Error processing files: {str(e)}")
        raise

def process_files(docx_path: str, excel_path: str, timeout: int = None) -> List[Tuple[str, str, float]]:
    """Process Word and Excel files to find and update matching content."""
    try:
        # Load input files
        print("Loading files...")
        excel_data = excel_to_list(excel_path)
        sentences = read_docx(docx_path)
        
        if not sentences:
            raise ValueError("No text found in the Word document")
            
        # Initialize LLM provider
        config = Config()
        llm_config = config.get_model_config()["llm"]
        llm = get_llm_provider(llm_config["provider"])
        
        # Process each sentence
        results = []
        print("\nProcessing sentences...")
        for sentence in tqdm(sentences.values(), desc="Analyzing", unit="sentence"):
            # Find exact matches in sentence
            matches = find_exact_matches(sentence, excel_data)
            if not matches:
                continue
            
            # Use LLM to update sentence with new values
            prompt = f"""Update this sentence with the new values from Excel:

Sentence: {sentence}

Excel data:
{chr(10).join(f'- {header}: {values}' for header, values in matches)}

Rules:
1. Keep the same sentence structure
2. Only update the numerical values
3. Keep all formatting (e.g., "$", "billion")
4. Do not change any other text

Output only the updated sentence."""

            try:
                updated_text = llm.analyze_text(prompt, timeout)
                if isinstance(updated_text, dict):
                    updated_text = updated_text.get("analysis", sentence)
                
                if updated_text and updated_text != sentence:
                    # Calculate confidence for each match
                    match_confidences = [
                        calculate_confidence(sentence, header, values, updated_text, len(matches))
                        for header, values in matches
                    ]
                    # Use average confidence if multiple matches
                    confidence = sum(match_confidences) / len(match_confidences)
                    
                    results.append((
                        sentence,
                        updated_text,
                        confidence
                    ))
                    
            except Exception as e:
                logger.error(f"Error updating sentence: {str(e)}")
                continue
        
        return results
            
    except Exception as e:
        logger.error(f"Error processing files: {str(e)}")
        raise

def update_document(docx_path: str, results: List[Tuple[str, str, float]], output_path: str) -> int:
    """Update the document with modified sentences and save to a new file."""
    try:
        doc = Document(docx_path)
        updates = {result[0].strip(): result[1].strip() for result in results if isinstance(result, tuple)}
        
        changes_made = _update_paragraphs(doc, updates)
        doc.save(output_path)
        
        print(f"\nMade {changes_made} changes")
        print(f"Updated document saved to: {output_path}")
        
        return changes_made
        
    except Exception as e:
        logger.error(f"Error updating document: {str(e)}")
        raise

def _update_paragraphs(doc: Document, updates: Dict[str, str]) -> int:
    """Update paragraphs in the document with new text while preserving formatting."""
    changes_made = 0
    
    for paragraph in doc.paragraphs:
        original_text = paragraph.text.strip()
        for orig_text, mod_text in updates.items():
            if orig_text in original_text:
                runs_info = _store_formatting(paragraph)
                _update_paragraph_text(paragraph, original_text, orig_text, mod_text, runs_info)
                changes_made += 1
                break
    
    return changes_made

def find_changes(updates: List[Tuple[str, str, float]]) -> int:
    """Count the number of actual changes in the processed results."""
    changes_made = 0

    for orig_text, mod_text, conf in updates:
        if orig_text != mod_text and conf > 0.1:  # Only count if text has changed and confidence is above 10%
            changes_made += 1
    
    return changes_made

def _store_formatting(paragraph) -> List[Dict]:
    """Store formatting information from paragraph runs."""
    runs_info = []
    if paragraph.runs:
        for run in paragraph.runs:
            runs_info.append({
                'text': run.text,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.font.bold,
                'italic': run.font.italic,
                'underline': run.font.underline,
                'color': run.font.color and run.font.color.rgb,
                'highlight_color': run.font.highlight_color,
            })
    return runs_info

def _update_paragraph_text(
    paragraph,
    original_text: str,
    orig_text: str,
    mod_text: str,
    runs_info: List[Dict]
) -> None:
    """Update paragraph text while preserving formatting."""
    # Clear paragraph content
    p = paragraph._p
    p.clear_content()
    
    # Replace text
    new_text = original_text.replace(orig_text, mod_text)
    
    if runs_info:
        _apply_formatted_runs(paragraph, new_text, runs_info)
    else:
        paragraph.add_run(new_text)

def _apply_formatted_runs(paragraph, new_text: str, runs_info: List[Dict]) -> None:
    """Apply formatted runs to the paragraph."""
    total_length = sum(len(run_info['text']) for run_info in runs_info)
    ratios = [len(run_info['text']) / total_length for run_info in runs_info]
    
    current_pos = 0
    for i, ratio in enumerate(ratios):
        length = int(len(new_text) * ratio) if i < len(ratios) - 1 else len(new_text) - current_pos
        run_text = new_text[current_pos:current_pos + length]
        current_pos += length
        
        new_run = paragraph.add_run(run_text)
        _apply_run_formatting(new_run, runs_info[i])

def _apply_run_formatting(run, formatting: Dict) -> None:
    """Apply stored formatting to a run."""
    if formatting['font_name']:
        run.font.name = formatting['font_name']
    if formatting['font_size']:
        run.font.size = formatting['font_size']
    if formatting['bold']:
        run.font.bold = formatting['bold']
    if formatting['italic']:
        run.font.italic = formatting['italic']
    if formatting['underline']:
        run.font.underline = formatting['underline']
    if formatting['color']:
        run.font.color.rgb = formatting['color']
    if formatting['highlight_color']:
        run.font.highlight_color = formatting['highlight_color']